"""
Cover Letter Generator — JD mode only, triggered by --cover-letter flag.

Pipeline:
  1. DuckDuckGo search → fetch company research (free, no API key)
  2. LLM generates a 4-paragraph cover letter as plain text
  3. Renders the letter in the terminal via Rich
  4. Writes a professionally formatted .docx file to the current directory

The .docx follows a formal business letter layout:
  - Candidate contact block (top left)
  - Date
  - Hiring team / company block
  - Salutation
  - Four body paragraphs (LLM output)
  - Professional sign-off
"""

import re
from datetime import datetime
from pathlib import Path
from rich.console import Console
from rich.rule import Rule
from rich.panel import Panel

from rich.prompt import Prompt

from pipeline import llm_client
from prompts import cover_letter_prompt

console = Console()


# ── Step 0: Company website lookup + confirmation ─────────────────────────────

def _find_official_website(company_name: str) -> str | None:
    """
    Searches DuckDuckGo for the company's official website.
    Returns the first result URL or None if nothing is found.
    """
    try:
        from duckduckgo_search import DDGS
        with DDGS(timeout=10) as ddgs:
            results = ddgs.text(f"{company_name} official website", max_results=3)
            for r in results:
                url = r.get("href", "").strip()
                if url:
                    return url
    except Exception:
        pass
    return None


def _confirm_or_correct_url(company_name: str, found_url: str | None) -> str | None:
    """
    Shows the found URL (or a not-found notice) and asks the user to confirm.

    - If confirmed       → returns the found URL.
    - If rejected        → prompts the user for the correct URL and returns it.
    - If nothing found   → prompts the user for the URL and returns it.
    - If user skips URL  → returns None (research will proceed on name alone).
    """
    console.print()

    if found_url:
        console.print(
            f"  [dim]Company:[/dim]  [bold]{company_name}[/bold]\n"
            f"  [dim]Website:[/dim]  [cyan]{found_url}[/cyan]"
        )
    else:
        console.print(
            f"  [dim]Company:[/dim]  [bold]{company_name}[/bold]\n"
            f"  [yellow]No official website found via search.[/yellow]"
        )

    console.print()
    choice = Prompt.ask(
        "  [bold]Is this the correct company?[/bold] "
        "[[green]y[/green]/[red]n[/red]]",
        choices=["y", "n"],
        default="y",
        show_choices=False,
    )

    if choice == "y":
        console.print()
        return found_url

    # User rejected or no URL found — ask for the correct one
    console.print()
    correct_url = Prompt.ask(
        "  [bold]Enter the correct official URL[/bold] "
        "[dim](leave blank to proceed without a URL)[/dim]"
    ).strip()

    console.print()
    return correct_url if correct_url else None


# ── Step 1: Company research ──────────────────────────────────────────────────

def _research_company(company_name: str, website_url: str | None = None) -> str:
    """
    Uses DuckDuckGo to fetch a concise company summary for the LLM prompt.
    If a confirmed website URL is provided, a site-specific query is added
    to anchor results to the correct company rather than a namesake.
    Falls back gracefully if search fails.
    """
    try:
        from duckduckgo_search import DDGS

        queries = [
            f"{company_name} company product what they do",
            f"{company_name} engineering culture mission",
        ]

        # Anchor results to the confirmed domain when available
        if website_url:
            from urllib.parse import urlparse
            domain = urlparse(website_url).netloc.lstrip("www.")
            if domain:
                queries.insert(0, f"site:{domain} about engineering")

        snippets: list[str] = []
        with DDGS(timeout=10) as ddgs:
            for query in queries:
                results = ddgs.text(query, max_results=3)
                for r in results:
                    body = r.get("body", "").strip()
                    if body:
                        snippets.append(body)

        if not snippets:
            return f"Company: {company_name}. No additional research available."

        return " ".join(snippets)[:1500]

    except Exception:
        return (
            f"Company: {company_name}. "
            "Research unavailable — infer context from the job description."
        )


# ── Step 2: LLM cover letter generation ──────────────────────────────────────

def _generate_letter(
    resume: dict,
    jd_text: str,
    jd_context: dict,
    company_name: str,
    company_research: str,
) -> str:
    """Calls the LLM and returns the cover letter body as a plain text string."""
    user_prompt = cover_letter_prompt.build(
        resume=resume,
        jd_text=jd_text,
        jd_context=jd_context,
        company_name=company_name,
        company_research=company_research,
    )

    return llm_client.chat(
        cover_letter_prompt.SYSTEM,
        user_prompt,
        temperature=0.5,    # Slightly higher — cover letters benefit from natural prose variation
        max_tokens=1024,    # 4 tight paragraphs never need more than this
    )


# ── Step 3: Terminal render ───────────────────────────────────────────────────

def _render_letter(
    letter_body: str,
    candidate_name: str,
    company_name: str,
    contact: dict,
) -> None:
    """Renders the cover letter in the terminal using Rich."""
    console.print()
    console.print(Rule("[bold cyan]COVER LETTER[/bold cyan]", style="cyan"))
    console.print()

    date_str = datetime.now().strftime("%B %d, %Y")

    header = (
        f"[bold]{candidate_name}[/bold]\n"
        f"[dim]{contact.get('email', '')}  ·  {contact.get('mobile', '')}[/dim]\n"
        f"[dim]{contact.get('github', '')}  ·  {contact.get('linkedin', '')}[/dim]\n\n"
        f"[dim]{date_str}[/dim]\n\n"
        f"Hiring Team\n"
        f"{company_name}\n"
    )
    console.print(Panel(header, border_style="dim", padding=(0, 2)))
    console.print()
    console.print("  [dim]Dear Hiring Manager,[/dim]\n")

    for paragraph in letter_body.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            console.print(f"  {paragraph}\n")

    console.print(f"  [dim]Sincerely,[/dim]")
    console.print(f"  [bold]{candidate_name}[/bold]")
    console.print()
    console.print(Rule(style="cyan"))


# ── Step 4: Write .docx ───────────────────────────────────────────────────────

def _write_docx(
    letter_body: str,
    candidate_name: str,
    company_name: str,
    contact: dict,
) -> Path:
    """
    Writes the cover letter as a properly formatted .docx business letter.

    Layout:
      - Candidate contact block (name bold, details normal)
      - Blank line
      - Date
      - Blank line
      - Recipient block (Hiring Team + Company)
      - Blank line
      - Salutation
      - Blank line
      - Four body paragraphs (justified, first-line indent)
      - Blank line
      - Sign-off + candidate name

    Page: US Letter, 1.25-inch side margins (standard business letter),
    1-inch top/bottom, Arial 11pt, 1.15 line spacing.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)   # Standard business letter margins
    section.right_margin  = Inches(1.25)

    # ── Default font ──────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = Pt(13.8)   # ~1.15 line spacing

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _p(text: str = "", bold: bool = False, italic: bool = False,
           size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT,
           space_before: int = 0, space_after: int = 6,
           first_line_indent: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        if text:
            run = p.add_run(text)
            run.bold        = bold
            run.italic      = italic
            run.font.size   = Pt(size)
            run.font.name   = "Arial"

    def _blank(space: int = 8) -> None:
        _p(space_after=space)

    # ── Candidate contact block ───────────────────────────────────────────────
    _p(candidate_name, bold=True, size=13, space_after=2)
    if contact.get("email"):
        _p(contact["email"], size=10, space_after=1)
    if contact.get("mobile"):
        _p(contact["mobile"], size=10, space_after=1)
    if contact.get("github"):
        _p(contact["github"], size=10, space_after=1)
    if contact.get("linkedin"):
        _p(contact["linkedin"], size=10, space_after=1)

    _blank(12)

    # ── Date ─────────────────────────────────────────────────────────────────
    _p(datetime.now().strftime("%B %d, %Y"), space_after=12)

    # ── Recipient block ───────────────────────────────────────────────────────
    _p("Hiring Team", space_after=2)
    _p(company_name, space_after=12)

    # ── Salutation ────────────────────────────────────────────────────────────
    _p("Dear Hiring Manager,", space_after=12)

    # ── Body paragraphs (LLM output) ──────────────────────────────────────────
    paragraphs = [p.strip() for p in letter_body.split("\n\n") if p.strip()]
    for para in paragraphs:
        _p(
            para,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=True,
            space_after=10,
        )

    _blank(12)

    # ── Sign-off ──────────────────────────────────────────────────────────────
    _p("Sincerely,", space_after=2)
    _blank(18)                          # Space for handwritten signature
    _p(candidate_name, bold=True, space_after=1)
    if contact.get("email"):
        _p(contact["email"], size=10, space_after=1)

    # ── Save ─────────────────────────────────────────────────────────────────
    def _slug(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", "_", s.lower()).strip("_")

    year     = datetime.now().year
    filename = f"cover_letter_{_slug(candidate_name)}_{_slug(company_name)}_{year}.docx"
    path     = Path.cwd() / filename
    doc.save(str(path))
    return path


# ── Public entry point ────────────────────────────────────────────────────────

def generate_cover_letter(
    resume: dict,
    jd_text: str,
    jd_context: dict,
    company_name: str,
) -> None:
    """
    Orchestrates all cover letter generation steps:
      website lookup → user confirmation → research → generate → render → save as .docx

    Raises CompanyMismatchError if the user rejects the found company website.

    Args:
        resume:       The optimised resume dict (post Pass 1/2 + approval).
        jd_text:      Raw job description text string.
        jd_context:   Parsed JD context from jd_parser.parse_jd().
        company_name: Company name string from --company-name flag.
    """
    contact        = resume.get("contact", {})
    candidate_name = resume.get("name", "Candidate")

    # Step 0: Website lookup + user confirmation / correction
    console.print("[cyan]→[/cyan] Looking up company website...")
    found_url    = _find_official_website(company_name)
    website_url  = _confirm_or_correct_url(company_name, found_url)

    # Step 1: Research — anchored to confirmed URL when available
    console.print("[cyan]→[/cyan] Researching company...")
    company_research = _research_company(company_name, website_url)
    console.print("[green]✓[/green] Company research complete.\n")

    # Step 2: Generate
    console.print("[cyan]→[/cyan] Generating cover letter...")
    letter_body = _generate_letter(
        resume, jd_text, jd_context, company_name, company_research
    )
    console.print("[green]✓[/green] Cover letter generated.\n")

    # Step 3: Render in terminal
    _render_letter(letter_body, candidate_name, company_name, contact)

    # Step 4: Save as .docx
    docx_path = _write_docx(letter_body, candidate_name, company_name, contact)
    console.print(
        f"[bold green]✓ Cover letter saved to:[/bold green] [cyan]{docx_path}[/cyan]\n"
    )