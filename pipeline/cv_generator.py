"""
CV Generator — JD mode only, triggered by --cv flag.

Pipeline:
  1. DuckDuckGo search → fetch company research (free, no API key)
  2. LLM generates full CV as plain text using resume + JD + company research
  3. Renders the CV in the terminal via Rich
  4. Writes a professionally formatted .docx file to the project assets directory
"""

import re
from pathlib import Path
from openai import OpenAI
from rich.console import Console
from rich.rule import Rule

from config import settings
from prompts import cv_prompt

client  = OpenAI(
    base_url="https://models.github.ai/inference",
    api_key=settings.github_access_token,
)
console = Console()

def _research_company(company_name: str) -> str:
    """
    Uses DuckDuckGo to search for the company and returns a concise
    research summary as a plain string for the LLM prompt.
    Falls back gracefully if search fails.
    """
    try:
        from duckduckgo_search import DDGS

        queries = [
            f"{company_name} company product what they do",
            f"{company_name} engineering culture tech stack",
        ]

        snippets: list[str] = []
        with DDGS() as ddgs:
            for query in queries:
                results = ddgs.text(query, max_results=3)
                for r in results:
                    body = r.get("body", "").strip()
                    if body:
                        snippets.append(body)

        if not snippets:
            return f"Company: {company_name}. No additional research available."

        combined = " ".join(snippets)
        return combined[:1500]

    except Exception:
        return (
            f"Company: {company_name}. "
            "Research unavailable — infer context from the job description."
        )

def _generate_cv(
    resume: dict,
    jd_text: str,
    jd_context: dict,
    company_name: str,
    company_research: str,
) -> str:
    """Calls the LLM and returns the CV as a plain text string."""
    user_prompt = cv_prompt.build(
        resume=resume,
        jd_text=jd_text,
        jd_context=jd_context,
        company_name=company_name,
        company_research=company_research,
    )

    response = client.chat.completions.create(
        model=settings.model,
        temperature=0.4,
        max_tokens=3000,
        messages=[
            {"role": "system", "content": cv_prompt.SYSTEM},
            {"role": "user",   "content": user_prompt},
        ],
    )

    return response.choices[0].message.content.strip()

def _render_cv(cv_text: str) -> None:
    """Renders the plain text CV in the terminal using Rich formatting."""
    console.print()
    console.print(Rule("[bold cyan]GENERATED CV[/bold cyan]", style="cyan"))
    console.print()

    section_headers = {
        "PROFESSIONAL SUMMARY",
        "CORE COMPETENCIES",
        "PROFESSIONAL EXPERIENCE",
        "PROJECTS",
        "EDUCATION",
    }

    lines = cv_text.splitlines()
    for i, line in enumerate(lines):
        stripped = line.strip()

        if i == 0:
            console.print(f"\n[bold white]{stripped}[/bold white]")
        elif i == 1:
            console.print(f"[dim]{stripped}[/dim]")
        elif stripped in section_headers:
            console.print()
            console.print(Rule(f"[bold yellow]{stripped}[/bold yellow]", style="yellow"))
        elif stripped.startswith("•"):
            console.print(f"  [cyan]•[/cyan] {stripped[1:].strip()}")
        elif stripped.startswith("Tech:"):
            console.print(f"  [dim]{stripped}[/dim]")
        elif stripped == "":
            console.print()
        else:
            console.print(f"  {stripped}")

    console.print()
    console.print(Rule(style="cyan"))

def _assets_dir() -> Path:
    assets = Path(__file__).resolve().parents[1] / "assets"
    assets.mkdir(parents=True, exist_ok=True)
    return assets


def _safe_filename_part(value: str) -> str:
    cleaned = re.sub(r"[<>:\"/\\|?*\x00-\x1F]", "_", value).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")
    return cleaned or "resume"


def _write_docx(cv_text: str, resume_name: str) -> Path:
    """
    Writes the CV as a professionally formatted .docx file.

    Structure mirrors the plain text output:
      - Candidate name → large bold heading
      - Contact line   → smaller centered line
      - Section headers → Heading 1 style with bottom border
      - Role / project headers → Heading 2 style
      - Body paragraphs → Normal style, justified
      - Bullet lines   → native Word list bullets (never unicode)
      - Tech lines     → italic Normal

    Page: US Letter, 1-inch margins, Arial 11pt body.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    section = doc.sections[0]
    section.page_width  = int(8.5 * 914400 / 100 * 100)
    section.page_height = int(11  * 914400 / 100 * 100)
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def _add_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_name(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold      = True
        run.font.size = Pt(20)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_after = Pt(2)

    def _add_contact(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size  = Pt(9)
        run.font.name  = "Arial"
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        p.paragraph_format.space_after = Pt(8)

    def _add_section_header(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold       = True
        run.font.size  = Pt(11)
        run.font.name  = "Arial"
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
        _add_bottom_border(p)

    def _add_role_header(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold      = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

    def _add_body(text: str, italic: bool = False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size   = Pt(11)
        run.font.name   = "Arial"
        run.font.italic = italic
        p.paragraph_format.space_after = Pt(2)

    def _add_bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Arial"
        p.paragraph_format.space_after = Pt(1)

    SECTION_HEADERS = {
        "PROFESSIONAL SUMMARY",
        "CORE COMPETENCIES",
        "PROFESSIONAL EXPERIENCE",
        "PROJECTS",
        "EDUCATION",
    }

    lines = cv_text.splitlines()
    for i, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            continue

        if i == 0:
            _add_name(stripped)

        elif i == 1:
            _add_contact(stripped)

        elif stripped in SECTION_HEADERS:
            _add_section_header(stripped)

        elif stripped.startswith("•"):
            _add_bullet(stripped[1:].strip())

        elif stripped.startswith("Tech:"):
            _add_body(stripped, italic=True)

        else:
            is_sub_header = (
                len(stripped) < 100
                and not stripped.endswith((".", ",", ";", ":"))
                and ("·" in stripped or "–" in stripped or "—" in stripped)
            )
            if is_sub_header:
                _add_role_header(stripped)
            else:
                _add_body(stripped)

    filename = f"{_safe_filename_part(resume_name)}_CV.docx"
    path = _assets_dir() / filename
    doc.save(str(path))
    return path

def generate_cv(
    resume: dict,
    jd_text: str,
    jd_context: dict,
    company_name: str,
    resume_name: str,
) -> None:
    """
    Orchestrates all CV generation steps:
      research → generate → render in terminal → save as .docx

    Args:
        resume:       The optimised resume dict (post Pass 1/2 + approval).
        jd_text:      Raw job description text string.
        jd_context:   Parsed JD context from jd_parser.parse_jd().
        company_name: Company name string from --company-name flag.
        resume_name:  Value used for naming the generated CV file.
    """
    # Step 1: Research
    console.print("[cyan]→[/cyan] Researching company...")
    company_research = _research_company(company_name)
    console.print("[green]✓[/green] Company research complete.\n")

    # Step 2: Generate
    console.print("[cyan]→[/cyan] Generating CV...")
    cv_text = _generate_cv(resume, jd_text, jd_context, company_name, company_research)
    console.print("[green]✓[/green] CV generated.\n")

    # Step 3: Render in terminal
    _render_cv(cv_text)

    # Step 4: Save as .docx
    docx_path = _write_docx(cv_text, resume_name)
    console.print(
        f"[bold green]✓ CV saved to:[/bold green] [cyan]{docx_path}[/cyan]\n"
    )
